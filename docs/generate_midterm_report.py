"""生成毕业设计中期报告 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_text(cell, text, bold=False, font_size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """设置单元格文本格式"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.bold = bold


def add_heading_styled(doc, text, level=1):
    """添加带格式的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        run.font.color.rgb = RGBColor(0, 0, 0)


def add_paragraph_styled(doc, text, bold=False, indent=False, font_size=12, first_line_indent=True):
    """添加带格式的段落"""
    p = doc.add_paragraph()
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.bold = bold
    return p


def add_bullet(doc, text, level=0, font_size=12):
    """添加项目符号段落"""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.5
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p


def main():
    doc = Document()

    # ---- 页面设置 ----
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ---- 封面标题 ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(60)
    title.paragraph_format.space_after = Pt(30)
    run = title.add_run("本科毕业设计（论文）中期检查报告")
    run.font.size = Pt(22)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.bold = True

    # ---- 基本信息 ----
    info_items = [
        ("课题名称", "基于智能体的网站和移动应用无障碍检测系统设计与实现"),
        ("学生姓名", "袁烁"),
        ("学    号", "162210417"),
        ("学    院", "计算机科学与技术学院/软件学院"),
        ("专    业", "人工智能创新班"),
        ("指导教师", "李丕绩 教授"),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 2.0
        run_label = p.add_run(f"{label}：")
        run_label.font.size = Pt(14)
        run_label.font.name = "宋体"
        run_label._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run_label.bold = True
        run_value = p.add_run(value)
        run_value.font.size = Pt(14)
        run_value.font.name = "宋体"
        run_value._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_before = Pt(40)
    run = date_p.add_run("2026 年 4 月")
    run.font.size = Pt(14)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # ---- 分页 ----
    doc.add_page_break()

    # ================================================================
    # 一、已完成的内容
    # ================================================================
    add_heading_styled(doc, "一、已完成的内容", level=1)

    # 1. 系统整体架构
    add_heading_styled(doc, "1. 系统整体架构设计与搭建", level=2)
    items = [
        "完成了前后端分离的整体架构设计：后端采用 Python + FastAPI 框架，前端采用 React + TypeScript + Ant Design 技术栈。",
        "完成 Docker Compose 容器化部署方案，包含 PostgreSQL（pgvector）、Redis、后端服务、前端服务四个容器。",
        "完成数据库设计，定义了 7 个核心数据模型：Project（项目）、Task（任务）、Page（页面）、Issue（问题）、Report（报告）、Annotation（标注）、AgentMemory（智能体记忆）。",
        "完成 RESTful API 层设计与实现，共 6 组路由（项目、任务、报告、问题、页面、标注），以及基于 WebSocket 的实时通信机制。",
    ]
    for item in items:
        add_bullet(doc, item)

    # 2. 智能体引擎
    add_heading_styled(doc, "2. 智能体引擎（核心模块一）", level=2)
    items = [
        "基于 LangGraph 框架实现了完整的多智能体编排状态机（Orchestrator），包含 Plan → Explore → Test → Report 四个节点的循环决策流程，支持迭代上限安全守卫机制。",
        "实现了四个专用智能体：MasterAgent（任务规划与决策）、CrawlerAgent（爬取策略建议）、DetectorAgent（AI 无障碍检测与多模态视觉分析）、ReporterAgent（报告生成与建议）。",
        "完成 LLM Provider 工厂模式设计，支持 OpenAI、Anthropic Claude、本地 Ollama 三种大语言模型后端的灵活切换。",
        "实现了双层记忆机制：短期记忆（基于有界队列，容量 100 条，支持按类型过滤）和长期记忆（PostgreSQL 持久化存储，支持按领域召回历史检测经验）。",
    ]
    for item in items:
        add_bullet(doc, item)

    # 3. 自适应遍历器
    add_heading_styled(doc, "3. 自适应遍历器（核心模块二）", level=2)
    items = [
        "实现了自适应爬虫（AdaptiveCrawler），基于 Playwright 浏览器自动化框架进行页面导航与交互。",
        "实现启发式 URL 优先级排序策略（HeuristicStrategy），根据页面类型（登录、表单、结账等）、页面深度、父页面特征、URL 长度等多因素综合打分排序。",
        "实现基于内容哈希和 DOM 结构签名的语义去重机制（SemanticDeduplicator），可自动识别模板化页面，在检测到 3 个同模板页面后跳过后续相似页面。",
        "实现应用状态图（ApplicationStateGraph），以有向图形式动态记录页面探索状态，支持 7 种节点状态和深度过滤查询。",
        "实现页面分析器（PageAnalyzer），可提取完整的无障碍树（A11y Tree）、标题结构、表单元素、链接、交互元素、地标区域等 DOM 结构信息，以及最多 150 个元素的边界框坐标。",
    ]
    for item in items:
        add_bullet(doc, item)

    # 4. 突发事件响应器
    add_heading_styled(doc, "4. 突发事件响应器（核心模块三）", level=2)
    items = [
        "实现了 EventResponder，可实时检测和自动处理四类突发界面事件：Cookie 同意横幅、弹窗/模态框、广告覆盖层、Alert 对话框。",
        "采用基于 CSS 选择器和 ARIA 属性的多策略检测方法，支持中英文双语按钮模式匹配（如 Accept/接受、Close/关闭），具备 Escape 键降级处理策略。",
        "事件检测与处理结果通过 WebSocket 实时推送至前端界面展示。",
    ]
    for item in items:
        add_bullet(doc, item)

    # 5. 人机协同工作台
    add_heading_styled(doc, "5. 人机协同工作台（核心模块四）", level=2)
    items = [
        "完成 Dashboard 首页，展示任务统计概览（总任务数、已完成、问题总数、已扫描页面数）和最近任务列表。",
        "完成项目管理页面，支持项目的创建、查看、删除等 CRUD 操作。",
        "完成任务创建页面，支持配置目标 URL、WCAG 检测级别（A/AA/AAA）、最大深度与页面数、视口尺寸、AI 检测与视觉检测等参数。",
        "完成任务详情页面，包含实时进度监控、问题列表、Agent 推理日志时间线、页面截图与标注四个 Tab 页，通过 WebSocket 实时更新。",
        "完成报告查看页面，包含分数仪表盘、严重度分布图表、问题摘要与改进建议、可导出为 HTML/PDF/JSON 格式。",
        "实现了带标注的截图查看组件（AnnotatedScreenshot），支持在截图上叠加显示问题边界框，以严重度颜色编码高亮，鼠标悬停显示问题详情。",
    ]
    for item in items:
        add_bullet(doc, item)

    # 6. 无障碍检测引擎
    add_heading_styled(doc, "6. 无障碍检测引擎", level=2)
    items = [
        "实现了规则检测引擎框架（DetectionEngine + RuleRegistry），采用注册表模式管理检测规则，支持按 WCAG 级别自动过滤。",
        "实现了 7 条 WCAG 2.1 Level A 检测规则：图片替代文本检测（1.1.1）、替代文本质量检测（1.1.1）、表单标签检测（1.3.1）、标题层级检测（1.3.1）、页面语言检测（3.1.1）、地标区域检测（1.3.1）、链接文本检测（2.4.4）。",
        "实现了 AI 辅助检测，由 DetectorAgent 基于 LLM 分析 A11y 树和 HTML 片段，发现规则难以覆盖的语义级无障碍问题。",
        "实现了多模态视觉检测，将页面截图与元素边界框信息发送至视觉大语言模型（GPT-4V/Claude），进行视觉层面的无障碍分析。",
        "实现了可视化标注系统（VisualAnnotator），在截图上自动绘制问题标注框、编号标签和图例，支持中文字体渲染。",
    ]
    for item in items:
        add_bullet(doc, item)

    # 7. 端到端数据流
    add_heading_styled(doc, "7. 端到端数据流打通", level=2)
    add_paragraph_styled(
        doc,
        "从前端创建任务 → 后端 API 接收 → 后台启动 Agent Pipeline → 多轮 Plan/Explore/Test 循环 → 生成报告的完整流程已全部打通。检测过程中的实时进度（已发现页面数、已测试数、问题数、当前 URL、Agent 推理日志）通过 WebSocket 从后端持续推送至前端展示。",
    )

    # ================================================================
    # 二、待完成的任务
    # ================================================================
    doc.add_page_break()
    add_heading_styled(doc, "二、待完成的任务", level=1)

    # 表格
    table = doc.add_table(rows=6, cols=4, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["序号", "任务内容", "优先级", "预计时间"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    tasks = [
        (
            "1",
            "扩充 WCAG 检测规则，补充 AA 级别规则（键盘可访问性 2.1.1、颜色对比度 1.4.3、焦点可见 2.4.7、跳过导航 2.4.1 等），提升规则覆盖率",
            "高",
            "1 周",
        ),
        (
            "2",
            "完成系统测试与验证：选择 2-3 个代表性网站进行端到端测试，计算遍历覆盖率提升率、突发事件处理成功率、缺陷检测准确率（Precision/Recall）等关键指标",
            "高",
            "1 周",
        ),
        (
            "3",
            "完成前端标注审核页面开发，实现问题逐条审核、误报标记、人工注释等\u201c人在回路\u201d闭环标注功能（后端 API 已就绪）",
            "中",
            "3 天",
        ),
        (
            "4",
            "系统优化：完善 Agent 的 LangChain Tool Calling 集成、长期记忆写入机制、Redis 缓存利用，提升检测效率",
            "中",
            "1 周",
        ),
        (
            "5",
            "撰写毕业论文，清晰阐述系统的设计思想、技术实现、实验验证过程和结论",
            "高",
            "2-3 周",
        ),
    ]
    for row_idx, (num, content, priority, time) in enumerate(tasks, start=1):
        set_cell_text(table.rows[row_idx].cells[0], num, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.rows[row_idx].cells[1], content)
        set_cell_text(table.rows[row_idx].cells[2], priority, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.rows[row_idx].cells[3], time, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 设置列宽
    for row in table.rows:
        row.cells[0].width = Cm(1.2)
        row.cells[1].width = Cm(10.5)
        row.cells[2].width = Cm(1.5)
        row.cells[3].width = Cm(1.8)

    # ================================================================
    # 三、存在的问题和解决办法
    # ================================================================
    doc.add_page_break()
    add_heading_styled(doc, "三、存在的问题和解决办法", level=1)

    # 问题 1
    add_heading_styled(doc, "问题1：WCAG 规则覆盖率有限", level=2)
    add_paragraph_styled(
        doc,
        "现状：目前系统仅实现了 7 条 WCAG 2.1 Level A 级别的检测规则，而 WCAG 2.1 标准共包含 78 条成功标准，当前规则覆盖率较低，尤其缺少 AA 级别的关键规则。",
    )
    add_paragraph_styled(
        doc,
        '解决办法：优先补充 AA 级别中影响面最大的检测规则，包括颜色对比度检测（1.4.3）、键盘可访问性检测（2.1.1）、焦点可见性检测（2.4.7）等。同时，充分利用 AI 检测能力（DetectorAgent）弥补规则覆盖的不足——大语言模型的语义理解能力可以发现许多规则难以形式化表达的无障碍问题，形成\u201c规则检测 + AI 检测\u201d的互补机制。',
    )

    # 问题 2
    add_heading_styled(doc, "问题2：移动应用检测能力缺失", level=2)
    add_paragraph_styled(
        doc,
        '现状：任务书中要求支持\u201c移动应用\u201d无障碍检测，但当前系统仅支持 Web 端检测，尚未集成 Appium 等移动端自动化测试框架。',
    )
    add_paragraph_styled(
        doc,
        "解决办法：考虑到项目时间限制，采取两方面策略。一是通过移动端视口模拟（系统已支持自定义 viewport 尺寸配置）和响应式页面检测，覆盖移动 Web 场景下的无障碍问题。二是在论文中设计移动端扩展架构方案，说明如何通过集成 Appium 框架、复用现有 Agent 编排和检测引擎来实现原生移动应用的无障碍检测，作为未来工作进行讨论。",
    )

    # 问题 3
    add_heading_styled(doc, "问题3：缺少量化实验验证", level=2)
    add_paragraph_styled(
        doc,
        "现状：系统核心功能已开发完成，但缺少系统性的对比实验和量化指标数据，包括任务书中要求的遍历覆盖率提升率、突发事件处理成功率、缺陷检测准确率等关键指标。",
    )
    add_paragraph_styled(
        doc,
        "解决办法：设计严谨的对比实验方案。选择 2-3 个不同类型的代表性网站（如政府服务网站、电商网站、新闻资讯网站），分别进行实验：（1）与传统深度优先遍历策略对比，测量自适应遍历器在相同时间内的页面覆盖率提升；（2）构造包含各类弹窗和广告的测试场景，统计突发事件处理成功率；（3）邀请人工专家对测试网站进行无障碍审计，以专家结果为基准计算系统检测的精确率和召回率。",
    )

    # 问题 4
    add_heading_styled(doc, "问题4：人机协同闭环功能未完善", level=2)
    add_paragraph_styled(
        doc,
        '现状：后端的标注 API（支持确认、误报标记、重分类、评论四种操作）已完整实现，但前端标注审核页面目前仍为占位符状态，\u201c人在回路\u201d的闭环优化机制尚未打通。',
    )
    add_paragraph_styled(
        doc,
        "解决办法：基于已完成的后端 API 和前端 Ant Design 组件库，快速开发标注审核页面。核心功能包括：问题逐条浏览与审核、一键标记误报、严重度重分类、添加人工注释。预计 3 天内可完成开发，因为后端接口和前端技术栈均已就绪。",
    )

    # 问题 5
    add_heading_styled(doc, "问题5：LLM 调用成本与延迟较高", level=2)
    add_paragraph_styled(
        doc,
        "现状：每次完整检测任务需要多轮大语言模型调用（MasterAgent 决策 + DetectorAgent 分析 + ReporterAgent 报告生成），当待检测页面数量较多时，API 调用成本和响应延迟显著增加。",
    )
    add_paragraph_styled(
        doc,
        "解决办法：引入分级检测策略——对结构简单、问题模式明确的页面仅使用规则引擎进行快速检测，仅对结构复杂或规则检测结果存疑的页面启用 AI 深度分析。同时，支持使用本地部署的 Ollama 模型（如 Qwen、Llama 等）替代商用 API，在可接受的质量损失范围内大幅降低调用成本。此外，利用语义去重机制减少对相似页面的重复检测，从源头降低 LLM 调用次数。",
    )

    # ---- 保存 ----
    output_path = "/Users/guangzhi/Desktop/W4Agent/docs/中期检查报告.docx"
    doc.save(output_path)
    print(f"报告已保存至: {output_path}")


if __name__ == "__main__":
    main()
